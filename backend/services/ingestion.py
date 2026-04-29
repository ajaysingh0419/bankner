"""
BankNER Document Ingestion Service
====================================
Handles text extraction from multiple document sources:

  LOCAL UPLOAD
  - PDF (text-based)         → PyMuPDF (fitz)
  - PDF (scanned/image)      → PyMuPDF + Tesseract OCR
  - PNG / JPG / TIFF images  → Tesseract OCR
  - DOCX (Word)              → python-docx
  - TXT / CSV                → plain decode

  REMOTE SOURCES
  - Public URL               → requests fetch + content-type detection
  - Google Drive (public)    → Drive export API (no auth for public files)
  - Dropbox (shared link)    → dl=1 redirect trick
  - Paste from URL           → generic fetch + PDF/text detection
"""

import io
import os
import re
import tempfile
import requests
from pathlib import Path
from typing import Tuple
from dataclasses import dataclass
from enum import Enum

# ── Optional imports with graceful fallback ──────────────────────────────────
try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    from PIL import Image
    import pytesseract
    OCR_OK = True
except ImportError:
    OCR_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


class IngestionSource(str, Enum):
    UPLOAD_PDF    = "upload_pdf"
    UPLOAD_IMAGE  = "upload_image"
    UPLOAD_DOCX   = "upload_docx"
    UPLOAD_TXT    = "upload_txt"
    URL           = "url"
    GOOGLE_DRIVE  = "google_drive"
    DROPBOX       = "dropbox"


@dataclass
class IngestionResult:
    text: str
    source: IngestionSource
    filename: str
    page_count: int = 1
    ocr_used: bool = False
    char_count: int = 0
    warning: str = ""

    def __post_init__(self):
        self.char_count = len(self.text)


# ── PDF Extraction ────────────────────────────────────────────────────────────

def extract_pdf(data: bytes, filename: str = "document.pdf") -> IngestionResult:
    if not PYMUPDF_OK:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")

    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    ocr_used = False

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()

        # If page has no extractable text, use OCR
        if len(text) < 20 and OCR_OK:
            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            text = pytesseract.image_to_string(img, config="--psm 6")
            ocr_used = True

        pages.append(text)

    full_text = "\n\n".join(pages)
    doc.close()

    return IngestionResult(
        text=full_text,
        source=IngestionSource.UPLOAD_PDF,
        filename=filename,
        page_count=len(pages),
        ocr_used=ocr_used,
        warning="OCR was used for scanned pages — accuracy may vary" if ocr_used else "",
    )


# ── Image OCR ─────────────────────────────────────────────────────────────────

def extract_image(data: bytes, filename: str = "image.png") -> IngestionResult:
    if not OCR_OK:
        raise RuntimeError("Pillow/pytesseract not installed.")

    img = Image.open(io.BytesIO(data))

    # Upscale small images for better OCR accuracy
    w, h = img.size
    if w < 1000:
        scale = 1000 / w
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    # Convert to greyscale for better OCR
    img = img.convert("L")

    text = pytesseract.image_to_string(img, config="--psm 6 --oem 3")

    return IngestionResult(
        text=text,
        source=IngestionSource.UPLOAD_IMAGE,
        filename=filename,
        ocr_used=True,
        warning="Text extracted via OCR — verify critical entities manually",
    )


# ── DOCX Extraction ───────────────────────────────────────────────────────────

def extract_docx(data: bytes, filename: str = "document.docx") -> IngestionResult:
    if not DOCX_OK:
        raise RuntimeError("python-docx not installed.")

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    text = "\n".join(paragraphs)

    return IngestionResult(
        text=text,
        source=IngestionSource.UPLOAD_DOCX,
        filename=filename,
    )


# ── Plain Text ────────────────────────────────────────────────────────────────

def extract_txt(data: bytes, filename: str = "document.txt") -> IngestionResult:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")

    return IngestionResult(
        text=text,
        source=IngestionSource.UPLOAD_TXT,
        filename=filename,
    )


# ── URL Fetch ─────────────────────────────────────────────────────────────────

def _normalise_drive_url(url: str) -> Tuple[str, IngestionSource]:
    """Convert Google Drive share URLs to direct download URLs."""
    # https://drive.google.com/file/d/FILE_ID/view
    m = re.search(r"drive\.google\.com/file/d/([a-zA-Z0-9_-]+)", url)
    if m:
        file_id = m.group(1)
        return f"https://drive.google.com/uc?export=download&id={file_id}", IngestionSource.GOOGLE_DRIVE

    # https://docs.google.com/document/d/DOC_ID/edit
    m = re.search(r"docs\.google\.com/document/d/([a-zA-Z0-9_-]+)", url)
    if m:
        doc_id = m.group(1)
        return f"https://docs.google.com/document/d/{doc_id}/export?format=txt", IngestionSource.GOOGLE_DRIVE

    return url, IngestionSource.URL


def _normalise_dropbox_url(url: str) -> Tuple[str, IngestionSource]:
    """Convert Dropbox share links to direct download links."""
    if "dropbox.com" in url:
        url = re.sub(r"\?dl=0", "?dl=1", url)
        if "?dl=1" not in url:
            url += "?dl=1"
        return url, IngestionSource.DROPBOX
    return url, IngestionSource.URL


def extract_from_url(url: str) -> IngestionResult:
    """Fetch a document from a public URL and extract text."""
    source = IngestionSource.URL

    # Normalise known cloud storage URLs
    if "drive.google.com" in url or "docs.google.com" in url:
        url, source = _normalise_drive_url(url)
    elif "dropbox.com" in url:
        url, source = _normalise_dropbox_url(url)

    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; BankNER/1.0; document processor)",
    }

    try:
        resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        resp.raise_for_status()
    except requests.exceptions.Timeout:
        raise ValueError("Request timed out after 15 seconds")
    except requests.exceptions.ConnectionError:
        raise ValueError(f"Could not connect to: {url}")
    except requests.exceptions.HTTPError as e:
        raise ValueError(f"HTTP {resp.status_code}: {e}")

    content_type = resp.headers.get("content-type", "").lower()
    data = resp.content
    filename = url.split("/")[-1].split("?")[0] or "remote_document"

    # Route by content type
    if "pdf" in content_type or filename.endswith(".pdf"):
        result = extract_pdf(data, filename)
    elif "image" in content_type or any(filename.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".tiff"]):
        result = extract_image(data, filename)
    elif "word" in content_type or "officedocument" in content_type or filename.endswith(".docx"):
        result = extract_docx(data, filename)
    else:
        # Try as plain text
        result = extract_txt(data, filename)

    result.source = source
    return result


# ── Main dispatcher ───────────────────────────────────────────────────────────

def ingest(data: bytes, filename: str) -> IngestionResult:
    """
    Route uploaded file bytes to the correct extractor based on filename extension.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(data, filename)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"):
        return extract_image(data, filename)
    elif ext == ".docx":
        return extract_docx(data, filename)
    elif ext in (".txt", ".csv", ".tsv", ".text"):
        return extract_txt(data, filename)
    else:
        # Try PDF first, fallback to text
        try:
            return extract_pdf(data, filename)
        except Exception:
            return extract_txt(data, filename)


def capabilities() -> dict:
    return {
        "pdf": PYMUPDF_OK,
        "ocr": OCR_OK,
        "docx": DOCX_OK,
        "url_fetch": True,
        "google_drive": True,
        "dropbox": True,
    }
